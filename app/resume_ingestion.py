"""Immutable, order-preserving DOCX ingestion for ResumeDocument imports.

This module is intentionally separate from the legacy plain-text parser.  Both
the saved text and the structured document are derived from one ordered stream.
"""
from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from itertools import count
from pathlib import Path

from app.resume_models import EducationItem, LanguageItem, PersonalInfo, ProjectItem, ResumeBullet, ResumeDocument, ResumeSection, SkillGroup, UnresolvedItem

_HEADINGS = {"education": {"education", "academic background", "academic experience"}, "skills": {"skills", "technical skills", "core skills", "competencies", "technical capabilities"}, "projects": {"projects", "project experience", "selected projects", "engineering projects", "technical projects"}, "languages": {"languages", "language skills"}, "competitions": {"engineering teams & competitions", "teams & competitions", "competitions"}, "summary": {"summary", "profile", "professional summary", "career summary"}}
_BULLET = re.compile(r"^\s*(?:[-*•●▪◦]|[^\w\s])\s+(.*)$")
_DATE = re.compile(r"(?:expected\s+)?(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)[^|/]*20\d{2}|expected\s+(?:spring|summer|fall|winter)?\s*20\d{2}", re.I)
_EMAIL = re.compile(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}")
_PHONE = re.compile(r"(?:\+?\d[\d().\s-]{7,}\d)")
_UPPERCASE_NAME = re.compile(r"^[A-Z][A-Z .'-]{2,}$")


@dataclass
class DocumentBlock:
    id: str
    order: int
    block_type: str
    text: str
    style_name: str = ""
    heading_level: int = 0
    is_bold: bool = False
    font_size: float = 0.0
    bullet_level: int = 0
    table_id: str = ""
    row_index: int = -1
    cell_index: int = -1
    cell_texts: list[str] | None = None
    source_ref: str = ""
    def to_dict(self) -> dict: return asdict(self)


def extract_docx_blocks(path: str | Path) -> list[DocumentBlock]:
    """Read body XML order; never iterate paragraphs and tables separately."""
    from docx import Document
    document = Document(str(path)); paragraphs = {id(p._p): p for p in document.paragraphs}; tables = {id(t._tbl): t for t in document.tables}
    blocks, order, table_number = [], 0, 0
    for child in document.element.body.iterchildren():
        paragraph, table = paragraphs.get(id(child)), tables.get(id(child))
        if paragraph is not None:
            text = paragraph.text.strip()
            if not text: continue
            runs = [run for run in paragraph.runs if run.text]; style = getattr(paragraph.style, "name", "") or ""; sizes = [run.font.size.pt for run in runs if run.font.size]
            blocks.append(DocumentBlock(f"block-{order}", order, "paragraph", text, style, int(style.casefold().startswith("heading")), bool(runs and all(run.bold for run in runs)), max(sizes, default=0.0), int(bool(_BULLET.match(text))), source_ref=f"paragraph:{order}")); order += 1
        elif table is not None:
            table_id = f"table-{table_number}"; table_number += 1
            for row_index, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if not any(cells): continue
                blocks.append(DocumentBlock(f"block-{order}", order, "table_row", " | ".join(cell.replace("\n", " / ") for cell in cells if cell), table_id=table_id, row_index=row_index, cell_texts=cells, source_ref=f"{table_id}:row:{row_index}")); order += 1
    return blocks


def normalized_source_text(blocks: list[DocumentBlock]) -> str:
    return "\n".join(block.text for block in sorted(blocks, key=lambda item: item.order) if block.text).strip()


def blocks_to_dict(blocks: list[DocumentBlock]) -> list[dict]: return [block.to_dict() for block in blocks]


def _heading(block: DocumentBlock) -> tuple[str | None, str]:
    clean = re.sub(r"\s+", " ", block.text.strip().strip(":")).casefold()
    for kind, labels in _HEADINGS.items():
        if clean in labels: return kind, block.text.strip().strip(":")
    if clean and len(clean) <= 62 and (block.is_bold or block.font_size >= 11 or block.style_name.casefold().startswith("heading")) and (block.text.isupper() or block.text.endswith(":")):
        return "custom", block.text.strip().strip(":")
    return None, ""


def _split_skills(value: str) -> list[str]:
    parts, current, depth = [], [], 0
    for char in value:
        depth += char in "(["; depth -= char in ")]" and depth > 0
        if char == "," and not depth:
            item = "".join(current).strip()
            if item: parts.append(item)
            current = []
        else: current.append(char)
    item = "".join(current).strip()
    if item: parts.append(item)
    return parts


def _personal(blocks: list[DocumentBlock]) -> PersonalInfo:
    lines = [block.text for block in blocks[:2]]; info = PersonalInfo(full_name=lines[0] if lines else "")
    contact = " | ".join(lines[1:]); email, phone = _EMAIL.search(contact), _PHONE.search(contact); info.email = email.group(0) if email else ""; info.phone = phone.group(0) if phone else ""
    pieces = [piece.strip() for piece in contact.split("|") if piece.strip()]
    leftovers = [piece for piece in pieces if piece not in {info.email, info.phone, "LinkedIn"}]
    if leftovers: info.location, info.other_lines = leftovers[0], leftovers[1:]
    return info


def _sections(blocks: list[DocumentBlock]):
    preamble, result, current, kind, title = [], [], [], None, ""
    for block in blocks:
        detected, detected_title = _heading(block)
        if kind is None and not preamble and _UPPERCASE_NAME.fullmatch(block.text) and 2 <= len(block.text.split()) <= 5: detected = None
        if detected:
            if kind is None: preamble.extend(current)
            else: result.append((kind, title, current))
            kind, title, current = detected, detected_title, []
        else: current.append(block)
    if kind is None: preamble.extend(current)
    else: result.append((kind, title, current))
    return preamble, result


def _education(raw: list[DocumentBlock], ids):
    entities, unresolved = [], []
    for block in raw:
        values = [value.strip() for value in (block.cell_texts or [block.text]) if value.strip()]; lines = [line.strip() for value in values for line in value.splitlines() if line.strip()]; joined = " | ".join(values)
        degree = next((value for value in lines if re.search(r"\b(bachelor|master|ph\.?d|diploma|degree)\b", value, re.I)), "")
        school_line = next((value for value in lines if re.search(r"\b(university|college|institute|school)\b", value, re.I)), "")
        institution = re.split(r"\s+-\s+(?=concentration\s*:)", school_line, 1, flags=re.I)[0].strip()
        concentration = (match.group(1).strip() if (match := re.search(r"concentration\s*:\s*([^|/]+)", joined, re.I)) else "")
        date = next((match.group(0).strip() for value in values if (match := _DATE.search(value))), "")
        location = next((value for value in lines if value != school_line and re.search(r"\b(?:ON|QC|BC|AB|Canada|Ottawa|Toronto|Montreal|Montréal|Quebec|Québec)\b", value, re.I)), "")
        if degree and institution: entities.append(EducationItem(f"education-{next(ids)}", institution=institution, degree=degree, concentration=concentration, location=location, date_text=date, source_refs=[block.id], parser_confidence=.98))
        elif block.text: unresolved.append(UnresolvedItem(f"unresolved-{next(ids)}", block.text, "education", [block.id], "Could not safely separate education degree and institution", ["education"], .45))
    return entities, unresolved


def _project_row(block: DocumentBlock, ids):
    cells = [cell for cell in (block.cell_texts or []) if cell.strip()]
    if not cells: return None
    lines = [line.strip() for line in cells[0].splitlines() if line.strip()]
    if len(lines) == 1 and " | " in lines[0]: lines = [part.strip() for part in lines[0].split(" | ") if part.strip()]
    date = next((value for value in cells[1:] if _DATE.search(value)), "")
    return ProjectItem(f"project-{next(ids)}", name=lines[0] if lines else "", subtitle=" | ".join(lines[1:]), date_text=date, source_refs=[block.id], parser_confidence=.98) if lines and date else None


def _apply_ai_resolutions(document: ResumeDocument, classifier, blocks: list[DocumentBlock], ids) -> None:
    """Map only high-confidence, source-verbatim AI classifications.

    The classifier never writes the persisted document. Its response is
    untrusted and may replace an unresolved block only when every value is
    present in exactly those source blocks.
    """
    if classifier is None or not document.unresolved_items:
        return
    by_id = {block.id: block for block in blocks}; remaining = []
    for item in document.unresolved_items:
        try: proposals = classifier.resolve(item.source_section, [by_id[ref] for ref in item.source_refs if ref in by_id])
        except Exception:
            remaining.append(item); continue
        accepted = False
        for proposal in proposals if isinstance(proposals, list) else []:
            if not isinstance(proposal, dict) or float(proposal.get("confidence", 0) or 0) < .85:
                continue
            refs = proposal.get("source_block_ids")
            if not isinstance(refs, list) or set(refs) != set(item.source_refs) or not set(refs).issubset(by_id):
                continue
            source = " ".join(by_id[ref].text for ref in refs).casefold()
            values = [str(value).strip() for key, value in proposal.items() if key not in {"type", "confidence", "source_block_ids"} and isinstance(value, str) and value.strip()]
            if not values or any(value.casefold() not in source for value in values):
                continue
            if item.source_section == "education" and proposal.get("type") == "education" and proposal.get("institution") and proposal.get("degree"):
                document.education.append(EducationItem(f"education-{next(ids)}", institution=str(proposal["institution"]), degree=str(proposal["degree"]), concentration=str(proposal.get("concentration") or ""), location=str(proposal.get("location") or ""), date_text=str(proposal.get("date_text") or ""), source_refs=list(refs), parser_confidence=float(proposal["confidence"]))); accepted = True
            elif item.source_section == "projects" and proposal.get("type") == "project" and proposal.get("name"):
                document.projects.append(ProjectItem(f"project-{next(ids)}", name=str(proposal["name"]), subtitle=str(proposal.get("subtitle") or ""), date_text=str(proposal.get("date_text") or ""), source_refs=list(refs), parser_confidence=float(proposal["confidence"]))); accepted = True
            if accepted: break
        if not accepted: remaining.append(item)
    document.unresolved_items = remaining


def parse_docx_resume(path: str | Path, semantic_classifier=None) -> tuple[ResumeDocument, list[DocumentBlock]]:
    blocks = extract_docx_blocks(path); preamble, sections = _sections(blocks); ids = count(1)
    document = ResumeDocument(personal_info=_personal(preamble), source_text=normalized_source_text(blocks), source_block_ids=[block.id for block in blocks], section_order=[], parser_version="ordered-semantic-v1")
    for kind, title, raw in sections:
        if kind == "education":
            document.section_order.append("education"); entities, unresolved = _education(raw, ids); document.education.extend(entities); document.unresolved_items.extend(unresolved)
        elif kind == "skills":
            document.section_order.append("skills")
            for block in raw:
                if re.match(r"languages?\s*:", block.text, re.I): document.languages.append(LanguageItem(f"language-{next(ids)}", block.text.split(":", 1)[1].strip(), source_refs=[block.id], parser_confidence=.99))
                elif ":" in block.text:
                    name, items = block.text.split(":", 1); document.skills.append(SkillGroup(f"skills-{next(ids)}", name.strip(), _split_skills(items), block.text, [block.id], .98))
                else: document.unresolved_items.append(UnresolvedItem(f"unresolved-{next(ids)}", block.text, "skills", [block.id], "Skill category is missing", ["skills", "custom"], .5))
            if document.languages: document.section_order.append("languages")
        elif kind == "projects":
            document.section_order.append("projects"); current = None
            for block in raw:
                candidate = _project_row(block, ids) if block.block_type == "table_row" else None
                if candidate: document.projects.append(candidate); current = candidate
                elif (match := _BULLET.match(block.text)) and current:
                    current.bullets.append(ResumeBullet(f"bullet-{next(ids)}", match.group(1).strip(), [block.id], .98)); current.source_refs.append(block.id)
                elif block.text: document.unresolved_items.append(UnresolvedItem(f"unresolved-{next(ids)}", block.text, "projects", [block.id], "Project content has no confirmed owner", ["project", "custom"], .5))
        elif kind in {"competitions", "custom"}:
            section = ResumeSection(f"section-{next(ids)}", "custom", title, [block.text for block in raw], True); document.custom_sections.append(section); document.section_order.append(f"custom:{section.id}")
        elif kind == "summary":
            # Summary is intentionally not part of the CareerOS document model
            # used for editing or generation. source_text still preserves the
            # imported words as candidate evidence.
            continue
        else:
            section = ResumeSection(f"section-{next(ids)}", kind, title, [block.text for block in raw], True); document.custom_sections.append(section); document.section_order.append(f"custom:{section.id}")
    _apply_ai_resolutions(document, semantic_classifier, blocks, ids)
    return document, blocks
