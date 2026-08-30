"""Conservative local parsing and rendering for ResumeDocument."""
from __future__ import annotations

import re
from dataclasses import dataclass
from itertools import count

from app.resume_models import EducationItem, ExperienceItem, LanguageItem, PersonalInfo, ProjectItem, ResumeBullet, ResumeDocument, ResumeSection, SkillGroup, UnresolvedItem


_HEADINGS = {
    "summary": {"summary", "profile", "professional summary", "career summary"},
    "experience": {"experience", "work experience", "professional experience", "employment", "work history"},
    "education": {"education", "academic background", "academic experience"},
    "skills": {"skills", "technical skills", "core skills", "competencies", "technical capabilities"},
    "projects": {"projects", "project experience", "selected projects", "engineering projects", "technical projects"},
    "certifications": {"certifications", "certificates", "licenses"},
    "languages": {"languages", "language skills"},
    "competitions": {"engineering teams & competitions", "teams & competitions", "competitions"},
}
_BULLET = re.compile(r"^\s*(?:[-*•●▪◦]|[^\w\s])\s+(.*)$")
_EMAIL = re.compile(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}")
_PHONE = re.compile(r"(?:\+?\d[\d().\s-]{7,}\d)")
_URL = re.compile(r"(?:https?://\S+|(?:linkedin\.com|github\.com|www\.)\S+)", re.I)
_UPPERCASE_NAME = re.compile(r"^[A-Z][A-Z .'-]{2,}$")


def _looks_like_uppercase_name(line: str) -> bool:
    """Avoid mistaking a top-of-resume all-caps name for a custom heading."""
    words = [word for word in re.split(r"\s+", line.strip()) if word]
    return bool(_UPPERCASE_NAME.fullmatch(line.strip()) and 2 <= len(words) <= 5)


def _heading(line: str) -> tuple[str | None, str]:
    clean = re.sub(r"\s+", " ", line.strip().strip(":")).casefold()
    for kind, labels in _HEADINGS.items():
        if clean in labels:
            return kind, line.strip().strip(":")
    if clean and len(clean) <= 62 and (line.strip().isupper() or line.strip().endswith(":")):
        return "custom", line.strip().strip(":")
    return None, ""


def _personal(lines: list[str]) -> PersonalInfo:
    info = PersonalInfo()
    if lines:
        info.full_name = lines[0].strip()
    rest = " | ".join(lines[1:])
    email = _EMAIL.search(rest); phone = _PHONE.search(rest)
    info.email = email.group(0) if email else ""; info.phone = phone.group(0) if phone else ""
    urls = _URL.findall(rest)
    for url in urls:
        lower = url.casefold()
        if "linkedin" in lower: info.linkedin = url
        elif "github" in lower: info.github = url
        elif not info.website: info.website = url
    # Keep only contact fragments not already promoted into structured fields.
    # Retaining the entire original line here would render the email/phone a
    # second time (for example, "email | Ottawa").
    residual: list[str] = []
    for line in lines[1:]:
        clean = line
        for known in (info.email, info.phone, *urls):
            if known:
                clean = clean.replace(known, "")
        for part in re.split(r"[|•]", clean):
            value = part.strip(" ,;-\t")
            if value and value not in residual:
                residual.append(value)
    if residual:
        info.location = residual[0]
        info.other_lines = residual[1:]
    return info


def _bullets(lines: list[str], ids) -> tuple[list[ResumeBullet], list[str]]:
    bullets, extra = [], []
    for line in lines:
        match = _BULLET.match(line)
        if match: bullets.append(ResumeBullet(f"bullet-{next(ids)}", match.group(1).strip()))
        elif line.strip(): extra.append(line.strip())
    return bullets, extra


def _split_skill_items(value: str) -> list[str]:
    """Split comma-separated skills while respecting nested parentheses."""
    result, current, depth = [], [], 0
    for char in value:
        if char in "([": depth += 1
        elif char in ")]" and depth: depth -= 1
        if char == "," and depth == 0:
            item = "".join(current).strip()
            if item: result.append(item)
            current = []
        else: current.append(char)
    item = "".join(current).strip()
    if item: result.append(item)
    return result


@dataclass
class DocumentBlock:
    id: str
    order: int
    text: str
    paragraph_style: str = ""
    is_bold: bool = False
    font_size: float = 0.0
    bullet_level: int = 0
    table_context: str = ""


def extract_docx_blocks(path) -> list[DocumentBlock]:
    """Read DOCX body order without flattening paragraphs and tables."""
    from docx import Document
    document = Document(str(path)); blocks: list[DocumentBlock] = []
    paragraphs = {id(paragraph._p): paragraph for paragraph in document.paragraphs}
    tables = {id(table._tbl): table for table in document.tables}
    order = 0
    for child in document.element.body.iterchildren():
        paragraph = paragraphs.get(id(child))
        table = tables.get(id(child))
        if paragraph is not None:
            text = paragraph.text.strip()
            if text:
                runs = [run for run in paragraph.runs if run.text]
                size = next((run.font.size.pt for run in runs if run.font.size), 0.0)
                blocks.append(DocumentBlock(f"paragraph-{order}", order, text, paragraph.style.name, bool(runs and all(run.bold for run in runs if run.text)), size, 1 if _BULLET.match(text) else 0)); order += 1
        elif table is not None:
            for row_index, row in enumerate(table.rows):
                text = " | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells if cell.text.strip())
                if text:
                    blocks.append(DocumentBlock(f"table-{len(blocks)}-row-{row_index}", order, text, table_context=f"table-{len(blocks)}")); order += 1
    return blocks


def _source_text(blocks: list[DocumentBlock]) -> str:
    return "\n".join(block.text for block in blocks)


def _education_from_blocks(blocks: list[DocumentBlock], ids) -> tuple[list[EducationItem], list[UnresolvedItem]]:
    values = [(block.text, block.id) for block in blocks if block.text]
    text = " | ".join(value for value, _ in values)
    institution = next((value for value, _ in values if re.search(r"\b(?:university|college|institute|school)\b", value, re.I)), "")
    degree = next((value for value, _ in values if re.search(r"\b(?:bachelor|master|ph\.?d|diploma|degree)\b", value, re.I)), "")
    concentration_match = re.search(r"concentration\s*:\s*([^|/]+)", text, re.I)
    date_match = re.search(r"(?:expected\s+)?(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)[^|/]*20\d{2}|expected\s+(?:spring|summer|fall|winter)?\s*20\d{2}", text, re.I)
    location = next((piece.strip() for value, _ in values for piece in re.split(r"[|/]", value) if re.search(r"\b(?:ON|QC|BC|AB|Canada|Ottawa|Toronto|Montreal)\b", piece, re.I)), "")
    if degree and institution:
        clean_degree = re.sub(r"\s*/\s*" + re.escape(institution) + r".*$", "", degree, flags=re.I).strip()
        return [EducationItem(f"education-{next(ids)}", institution=institution.split(" - ")[0].strip(), degree=clean_degree, concentration=concentration_match.group(1).strip() if concentration_match else "", location=location, date_text=date_match.group(0).strip() if date_match else "", source_refs=[ref for _, ref in values], parser_confidence=.96)], []
    unresolved = [UnresolvedItem(f"unresolved-{next(ids)}", text, "EDUCATION", [ref for _, ref in values], "Could not confidently identify both institution and degree", ["education"], .55)] if text else []
    return [], unresolved


def parse_docx_resume(path) -> ResumeDocument:
    """Hybrid deterministic import: preserve block order and leave ambiguity unresolved."""
    blocks = extract_docx_blocks(path); source = _source_text(blocks); ids = count(1)
    document = ResumeDocument(personal_info=_personal([block.text for block in blocks[:2]]), source_text=source, section_order=[], parser_version="hybrid-v1")
    sections: list[tuple[str, str, list[DocumentBlock]]] = []; current_kind, current_title, current_blocks = "preamble", "", []
    for block in blocks:
        kind, title = _heading(block.text)
        if current_kind == "preamble" and not current_blocks and _looks_like_uppercase_name(block.text): kind = None
        if kind:
            if current_kind != "preamble": sections.append((current_kind, current_title, current_blocks))
            current_kind, current_title, current_blocks = kind, title, []
        else: current_blocks.append(block)
    if current_kind != "preamble": sections.append((current_kind, current_title, current_blocks))
    for kind, title, raw in sections:
        if kind == "education":
            document.section_order.append("education"); entities, unresolved = _education_from_blocks(raw, ids); document.education.extend(entities); document.unresolved_items.extend(unresolved)
        elif kind == "skills":
            document.section_order.append("skills")
            for block in raw:
                if re.match(r"languages?\s*:", block.text, re.I):
                    language = re.split(r":", block.text, 1)[1].strip(); document.languages.append(LanguageItem(f"language-{next(ids)}", language=language, source_refs=[block.id], parser_confidence=.98)); continue
                if ":" in block.text:
                    name, items = block.text.split(":", 1); document.skills.append(SkillGroup(f"skills-{next(ids)}", name.strip(), _split_skill_items(items), block.text, [block.id], .97))
                elif block.text:
                    document.unresolved_items.append(UnresolvedItem(f"unresolved-{next(ids)}", block.text, title, [block.id], "Skill category is missing", ["skills", "custom"], .55))
            if document.languages and "languages" not in document.section_order: document.section_order.append("languages")
        elif kind == "projects":
            document.section_order.append("projects"); current = None
            for block in raw:
                if block.table_context:
                    left, _, right = block.text.partition(" | "); parts = [part.strip() for part in left.split("|") if part.strip()]
                    current = ProjectItem(f"project-{next(ids)}", name=parts[0] if parts else left, organization=parts[1] if len(parts) > 1 else "", date_text=right.strip(), source_refs=[block.id], parser_confidence=.94); document.projects.append(current)
                elif _BULLET.match(block.text) and current is not None:
                    current.bullets.append(ResumeBullet(f"bullet-{next(ids)}", _BULLET.match(block.text).group(1).strip())); current.source_refs.append(block.id)
                elif block.text:
                    document.unresolved_items.append(UnresolvedItem(f"unresolved-{next(ids)}", block.text, title, [block.id], "Project boundary is ambiguous", ["project", "custom"], .50))
        elif kind in {"competitions", "custom"}:
            section = ResumeSection(f"section-{next(ids)}", "custom", title, [block.text for block in raw], True); document.custom_sections.append(section); document.section_order.append(f"custom:{section.id}")
        elif kind == "summary":
            continue
        elif kind == "languages":
            document.section_order.append("languages"); document.languages.extend(LanguageItem(f"language-{next(ids)}", language=block.text, source_refs=[block.id], parser_confidence=.95) for block in raw if block.text)
        elif kind == "certifications": document.section_order.append("certifications"); document.certifications.extend(block.text for block in raw if block.text)
        elif kind == "experience":
            document.section_order.append("experience"); document.unresolved_items.extend(UnresolvedItem(f"unresolved-{next(ids)}", block.text, title, [block.id], "Experience boundary needs review", ["experience", "custom"], .50) for block in raw if block.text)
    return document


def parse_resume_text(text: str) -> ResumeDocument:
    """Parse only clear structure; source_text always retains every input line."""
    source = str(text or "").replace("\r", "")
    lines = [line.rstrip() for line in source.split("\n")]
    sections: list[tuple[str, str, list[str]]] = []
    preamble: list[str] = []; current_kind, current_title, current_lines = "preamble", "", []
    for line in lines:
        kind, title = _heading(line)
        # Names such as "XUAN NIE" are commonly all caps.  They occur before
        # the first section and must stay in personal_info, not become a custom
        # section merely because they are uppercase.
        if current_kind == "preamble" and not preamble and not current_lines and _looks_like_uppercase_name(line):
            kind, title = None, ""
        if kind:
            if current_kind == "preamble": preamble.extend(current_lines)
            else: sections.append((current_kind, current_title, current_lines))
            current_kind, current_title, current_lines = kind, title, []
        else:
            current_lines.append(line)
    if current_kind == "preamble": preamble.extend(current_lines)
    else: sections.append((current_kind, current_title, current_lines))
    document = ResumeDocument(personal_info=_personal([line.strip() for line in preamble if line.strip()]), source_text=source, section_order=[])
    ids = count(1)
    for kind, title, raw in sections:
        meaningful = [line.strip() for line in raw if line.strip()]
        if kind == "summary":
            continue
        elif kind == "skills":
            if "skills" not in document.section_order: document.section_order.append("skills")
            for line in meaningful:
                if re.match(r"languages?\s*:", line, re.I):
                    language = line.split(":", 1)[1].strip(); document.languages.append(LanguageItem(f"language-{next(ids)}", language=language, parser_confidence=.95));
                    if "languages" not in document.section_order: document.section_order.append("languages")
                else:
                    name, values = (line.split(":", 1) + [""])[:2] if ":" in line else (title or "Skills", line)
                    document.skills.append(SkillGroup(f"skills-{next(ids)}", name.strip(), _split_skill_items(values), line, parser_confidence=.85))
        elif kind == "certifications":
            if "certifications" not in document.section_order: document.section_order.append("certifications")
            document.certifications.extend(meaningful)
        elif kind == "languages":
            if "languages" not in document.section_order: document.section_order.append("languages")
            document.languages.extend(LanguageItem(f"language-{next(ids)}", language=line, parser_confidence=.9) for line in meaningful)
        elif kind == "experience":
            if "experience" not in document.section_order: document.section_order.append("experience")
            bullets, extra = _bullets(raw, ids); document.experience.append(ExperienceItem(f"experience-{next(ids)}", extra_lines=extra, bullets=bullets))
        elif kind == "education":
            if "education" not in document.section_order: document.section_order.append("education")
            document.education.append(EducationItem(f"education-{next(ids)}", extra_lines=meaningful))
        elif kind in {"projects", "competitions"}:
            if "projects" not in document.section_order: document.section_order.append("projects")
            bullets, extra = _bullets(raw, ids); document.projects.append(ProjectItem(f"project-{next(ids)}", name=title if kind == "competitions" else "", bullets=bullets, extra_lines=extra))
        else: document.custom_sections.append(ResumeSection(f"section-{next(ids)}", "custom", title or "Additional information", raw, True)); document.section_order.append(f"custom:{document.custom_sections[-1].id}")
    return document


def _contact_lines(info: PersonalInfo) -> list[str]:
    values = [info.email, info.phone, info.location, info.linkedin, info.github, info.website, *info.other_lines, *getattr(info, "other_links", [])]
    output: list[str] = []
    seen: set[str] = set()
    for raw in values:
        value = str(raw or "").strip()
        if not value:
            continue
        digits = re.sub(r"\D", "", value)
        key = f"phone:{digits[-10:]}" if len(digits) >= 10 and sum(ch.isdigit() for ch in value) >= 7 else re.sub(r"\s+", " ", value).casefold()
        if key in seen:
            continue
        seen.add(key)
        output.append(value)
    return [" | ".join(output)] if output else []


def resume_document_to_text(document: ResumeDocument, hidden_sections: list[str] | None = None) -> str:
    """Render a deterministic compatibility string from structured data."""
    out = [document.personal_info.full_name.strip(), *_contact_lines(document.personal_info)]
    hidden = set(hidden_sections or [])
    sections = {section.id: section for section in document.custom_sections}
    for key in document.section_order:
        if key in hidden: continue
        # Summary is retired product data. Legacy JSON can still be loaded, but
        # no Resume/CV text output emits it.
        if key == "summary": continue
        elif key == "experience" and document.experience:
            out.extend(["", "EXPERIENCE"])
            for item in document.experience:
                heading = " | ".join(value for value in (item.title, item.company, item.location, item.date_text) if value); out.extend(([heading] if heading else []) + item.extra_lines + ["- " + bullet.text for bullet in item.bullets])
        elif key == "education" and document.education:
            out.extend(["", "EDUCATION"])
            for item in document.education:
                heading = " | ".join(value for value in (item.institution, item.degree, item.field, item.location, item.date_text) if value); out.extend(([heading] if heading else []) + item.details + item.extra_lines)
        elif key == "skills" and document.skills:
            out.extend(["", "SKILLS"]); out.extend(group.raw_text or f"{group.name}: {', '.join(group.items)}" for group in document.skills)
        elif key == "projects" and document.projects:
            out.extend(["", "PROJECTS"])
            for item in document.projects: out.extend(([" | ".join(value for value in (item.name, item.subtitle, item.date_text) if value)] if any((item.name, item.subtitle, item.date_text)) else []) + item.extra_lines + ["- " + bullet.text for bullet in item.bullets])
        elif key == "certifications" and document.certifications: out.extend(["", "CERTIFICATIONS", *document.certifications])
        elif key == "languages" and document.languages: out.extend(["", "LANGUAGES", *(f"{item.language}{' - ' + item.proficiency if item.proficiency else ''}" if hasattr(item, 'language') else str(item) for item in document.languages)])
        elif key.startswith("custom:"):
            section = sections.get(key.split(":", 1)[1])
            if section and section.visible: out.extend(["", section.title, *section.raw_lines])
    return "\n".join(line for line in out if line is not None).strip()
