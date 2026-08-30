from __future__ import annotations

import re
import base64
import os
import subprocess
import time
import threading
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.resume_models import ResumeDocument


ACCENT = RGBColor(31, 78, 121)
BLACK = RGBColor(0, 0, 0)
BULLET = re.compile(r"^\s*(?:[-*•●▪◦]|[^\w\s])\s+")
_WORD_EXPORT_LOCK = threading.Lock()


def _commit_pdf(temp_target: Path, target: Path) -> bool:
    """Atomically publish a PDF with short retries for Windows viewer locks."""
    if not temp_target.is_file():
        return False
    target.parent.mkdir(parents=True, exist_ok=True)
    for delay in (0.0, 0.08, 0.18, 0.35, 0.6):
        if delay:
            time.sleep(delay)
        try:
            os.replace(temp_target, target)
            return target.is_file()
        except OSError:
            continue
    try:
        temp_target.unlink(missing_ok=True)
    except OSError:
        pass
    return False


def export_docx_pdf(source: Path, target: Path) -> bool:
    """Export through Word to a temporary PDF, then publish atomically.

    QPdfDocument can briefly retain a Windows handle to the old target.  Asking
    Word to overwrite that same file made export fail and caused CareerOS to
    fall back to the legacy flow renderer.  A unique temporary target avoids
    that failure; the final replace is retried after the viewer has released
    its handle.
    """
    if os.name != "nt" or source.suffix.casefold() != ".docx" or not source.is_file():
        return False
    target.parent.mkdir(parents=True, exist_ok=True)
    temp_target = target.with_name(f".{target.stem}-{os.getpid()}-{time.time_ns()}.tmp.pdf")
    source_text = str(source.resolve()).replace("'", "''")
    target_text = str(temp_target.resolve()).replace("'", "''")
    script = (
        "$ErrorActionPreference='Stop';$word=$null;$doc=$null;"
        f"try{{$word=New-Object -ComObject Word.Application;$word.Visible=$false;$word.DisplayAlerts=0;$doc=$word.Documents.Open('{source_text}',$false,$true);$doc.ExportAsFixedFormat('{target_text}',17)}}"
        "finally{if($doc){$doc.Close(0)};if($word){$word.Quit()}}"
    )
    encoded = base64.b64encode(script.encode("utf-16le")).decode("ascii")
    windows_dir = Path(os.environ.get("WINDIR", r"C:\Windows")).resolve()
    powershell = windows_dir / "System32" / "WindowsPowerShell" / "v1.0" / "powershell.exe"
    if not powershell.is_file():
        return False
    try:
        # Word is an STA COM server.  Serializing local automation prevents
        # startup preview, editor preview, and save from racing one another.
        with _WORD_EXPORT_LOCK:
            # The executable is an absolute System32 path and the encoded script
            # contains only resolved local paths, escaped for PowerShell literals.
            result = subprocess.run([str(powershell), "-NoProfile", "-NonInteractive", "-EncodedCommand", encoded], capture_output=True, timeout=75, creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0))  # nosec B603
        if result.returncode == 0:
            return _commit_pdf(temp_target, target)
        return False
    except (OSError, subprocess.SubprocessError):
        return False
    finally:
        try:
            temp_target.unlink(missing_ok=True)
        except OSError:
            pass


def _set_font(run, size: float, bold: bool = False, color=None, italic: bool = False, family: str = "Arial"):
    run.font.name = family
    run._element.rPr.rFonts.set(qn("w:ascii"), family)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), family)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "7")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "1F4E79")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def _heading(line: str) -> bool:
    value = line.strip().lstrip("#").strip().rstrip(":")
    return bool(value) and len(value) <= 60 and (value.isupper() or value.casefold() in {
        "education", "experience", "skills", "projects", "certifications",
        "work experience", "technical skills", "languages", "engineering teams & competitions",
    })


def _unique_contacts(info) -> list[str]:
    """Return user-facing contact values without duplicate phone variants."""
    values = [info.location, info.phone, info.email, info.linkedin, info.github, info.website, *info.other_lines, *getattr(info, "other_links", [])]
    out: list[str] = []
    seen: set[str] = set()
    for raw in values:
        value = str(raw or "").strip()
        if not value:
            continue
        digits = re.sub(r"\D", "", value)
        key = f"phone:{digits[-10:]}" if len(digits) >= 10 and sum(ch.isdigit() for ch in value) >= 7 else re.sub(r"\s+", " ", value).casefold()
        if key in seen:
            continue
        seen.add(key)
        out.append(value)
    return out


def _section_heading(document, title: str, family: str, body_size: float, before: float, after: float):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    _set_font(p.add_run(title), body_size + 1.2, True, ACCENT, family=family)
    _rule(p)
    return p


def _add_two_col_row(document, left_parts: list[tuple[str, bool, bool]], right: str, family: str, body_size: float, *, after: float = 0.0, left_ratio: float = .78):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    _remove_table_borders(table)
    usable = document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin
    left_ratio = max(.60, min(.90, float(left_ratio)))
    table.columns[0].width = int(usable * left_ratio)
    table.columns[1].width = int(usable * (1 - left_ratio))
    left, right_cell = table.rows[0].cells
    _set_cell_margins(left); _set_cell_margins(right_cell)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(after)
    lp.paragraph_format.line_spacing = 1.0
    for index, (text, bold, italic) in enumerate(left_parts):
        if not text:
            continue
        if index and lp.text:
            _set_font(lp.add_run(" | "), body_size, family=family)
        _set_font(lp.add_run(text), body_size, bold=bold, italic=italic, family=family)
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(after)
    if right:
        _set_font(rp.add_run(right), body_size, True, family=family)
    return table


def _add_bullet(document, text: str, family: str, body_size: float, after: float = .6):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(.14)
    p.paragraph_format.first_line_indent = Inches(-.10)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = .96
    _set_font(p.add_run("• "), body_size, True, ACCENT, family=family)
    _set_font(p.add_run(text), body_size, family=family)
    return p


def _split_activity_heading(line: str) -> tuple[str, str]:
    """Keep the original custom content, but pull an obvious trailing date right."""
    text = line.strip().lstrip("•").strip()
    match = re.search(r"\s*\|\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[^|]{2,40}(?:Present|\d{4}))\s*(?:/\s*(.+))?$", text, re.I)
    if not match:
        return text, ""
    left = text[:match.start()].strip(" |")
    date = match.group(1).strip()
    tail = (match.group(2) or "").strip()
    if tail:
        left = f"{left} | {tail}"
    return left, date


def render_structured_resume_docx(document_model: ResumeDocument, target: Path, options: dict | None = None, hidden_sections: list[str] | None = None, document_label: str = "") -> Path:
    """Render the canonical ResumeDocument with a compact, original-style one-page layout."""
    options = options or {}
    family = "Times New Roman" if str(options.get("font_family", "")).casefold().startswith("times") else "Arial"
    body_size = max(8.0, min(11.0, float(options.get("font_size", 8.7))))
    margin_top = float(options.get("margin_top", .48)); margin_bottom = float(options.get("margin_bottom", .48))
    margin_left = float(options.get("margin_left", .58)); margin_right = float(options.get("margin_right", .58))
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(margin_top); section.bottom_margin = Inches(margin_bottom)
    section.left_margin = Inches(margin_left); section.right_margin = Inches(margin_right)
    normal = document.styles["Normal"]
    normal.font.name = family; normal._element.rPr.rFonts.set(qn("w:ascii"), family); normal._element.rPr.rFonts.set(qn("w:hAnsi"), family)
    normal.font.size = Pt(body_size); normal.paragraph_format.space_after = Pt(0); normal.paragraph_format.line_spacing = .98

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER; name.paragraph_format.space_after = Pt(0)
    _set_font(name.add_run(document_model.personal_info.full_name.strip() or "Resume"), float(options.get("name_font_size", 17.5)), True, ACCENT, family=family)
    contacts = _unique_contacts(document_model.personal_info)
    if contacts:
        p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(1.2)
        _set_font(p.add_run("  |  ".join(contacts)), max(7.6, body_size - .4), family=family)
    if document_label:
        label = document.add_paragraph(); label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label.paragraph_format.space_before = Pt(1.5); label.paragraph_format.space_after = Pt(3.0)
        _set_font(label.add_run(document_label), max(body_size + 1.2, 10.5), True, ACCENT, family=family)

    hidden = set(hidden_sections or [])
    sections = {section.id: section for section in document_model.custom_sections}
    section_before = float(options.get("section_spacing_before", 4.2)); section_after = float(options.get("section_spacing_after", 1.1))

    for key in document_model.section_order:
        if key in hidden:
            continue
        if key == "education" and document_model.education:
            _section_heading(document, "EDUCATION", family, body_size, section_before, section_after)
            for item in document_model.education:
                first = [(item.degree or item.field or item.institution, True, False)]
                _add_two_col_row(document, first, item.date_text, family, body_size, after=0)
                second_parts: list[tuple[str, bool, bool]] = []
                if item.institution: second_parts.append((item.institution, False, False))
                concentration = item.concentration or (item.field if item.degree else "")
                if concentration: second_parts.append((f"Concentration: {concentration}" if not concentration.lower().startswith("concentration") else concentration, False, True))
                # Education concentration is a compact descriptor, not a new
                # paragraph. Give it enough width to stay on one line while the
                # short location remains right-aligned.
                _add_two_col_row(document, second_parts, item.location, family, max(7.6, body_size-.35), after=0, left_ratio=.86)
                for detail in [*item.details, *item.extra_lines]:
                    p = document.add_paragraph(); p.paragraph_format.space_after = Pt(.4); _set_font(p.add_run(detail), body_size, family=family)
        elif key == "experience" and document_model.experience:
            _section_heading(document, "EXPERIENCE", family, body_size, section_before, section_after)
            for item in document_model.experience:
                _add_two_col_row(document, [(item.title, True, False), (item.company, False, False)], item.date_text, family, body_size)
                if item.location:
                    p = document.add_paragraph(); p.paragraph_format.space_after = Pt(.3); _set_font(p.add_run(item.location), body_size-.1, italic=True, family=family)
                for line in item.extra_lines: _add_bullet(document, line, family, body_size)
                for bullet in item.bullets: _add_bullet(document, bullet.text, family, body_size)
        elif key == "skills" and document_model.skills:
            _section_heading(document, "TECHNICAL SKILLS", family, body_size, section_before, section_after)
            for group in document_model.skills:
                p = document.add_paragraph(); p.paragraph_format.space_after = Pt(.3); p.paragraph_format.line_spacing = .96
                if group.name:
                    _set_font(p.add_run(f"{group.name}: "), body_size, True, ACCENT, family=family)
                _set_font(p.add_run(", ".join(group.items) if group.items else group.raw_text), body_size, family=family)
        elif key == "languages" and document_model.languages:
            # The source resume places language immediately under technical skills.
            if "skills" not in document_model.section_order or document_model.section_order.index("skills") > document_model.section_order.index("languages"):
                _section_heading(document, "LANGUAGES", family, body_size, section_before, section_after)
            for item in document_model.languages:
                p = document.add_paragraph(); p.paragraph_format.space_after = Pt(.3)
                _set_font(p.add_run("Languages: "), body_size, True, ACCENT, family=family)
                text = item.language + (f" ({item.proficiency})" if item.proficiency else "")
                _set_font(p.add_run(text), body_size, family=family)
        elif key == "projects" and document_model.projects:
            _section_heading(document, "ENGINEERING PROJECTS", family, body_size, section_before, section_after)
            project_gap = max(2.2, min(4.5, section_before * .55))
            for project_index, item in enumerate(document_model.projects):
                left = [(item.name, True, False)]
                if item.subtitle: left.append((item.subtitle, False, False))
                _add_two_col_row(document, left, item.date_text, family, body_size)
                last_paragraph = None
                for line in item.extra_lines: last_paragraph = _add_bullet(document, line, family, body_size)
                for bullet in item.bullets: last_paragraph = _add_bullet(document, bullet.text, family, body_size)
                # Give each project a small visual breathing space while
                # keeping the professional one-page layout compact.
                if project_index < len(document_model.projects) - 1:
                    if last_paragraph is not None:
                        last_paragraph.paragraph_format.space_after = Pt(project_gap)
                    else:
                        spacer = document.add_paragraph(); spacer.paragraph_format.space_after = Pt(project_gap)
                        spacer.paragraph_format.line_spacing = .2
        elif key == "certifications" and document_model.certifications:
            _section_heading(document, "CERTIFICATIONS", family, body_size, section_before, section_after)
            for value in document_model.certifications:
                p = document.add_paragraph(); p.paragraph_format.space_after = Pt(.3); _set_font(p.add_run(str(value)), body_size, family=family)
        elif key.startswith("custom:"):
            section_model = sections.get(key.split(":", 1)[1])
            if not section_model or not section_model.visible:
                continue
            _section_heading(document, section_model.title, family, body_size, section_before, section_after)
            activity_gap = max(2.2, min(4.5, section_before * .55))
            last_activity_paragraph = None
            activity_started = False
            for raw in section_model.raw_lines:
                if BULLET.match(raw):
                    last_activity_paragraph = _add_bullet(document, BULLET.sub("", raw), family, body_size)
                else:
                    left, date = _split_activity_heading(raw)
                    # A non-bullet line starts the next team/activity. Give the
                    # previous entry the same breathing room used between Projects.
                    if activity_started and last_activity_paragraph is not None:
                        last_activity_paragraph.paragraph_format.space_after = Pt(activity_gap)
                    activity_started = True
                    if date:
                        parts = [(part.strip(), i == 0, False) for i, part in enumerate(left.split(" | ")) if part.strip()]
                        _add_two_col_row(document, parts, date, family, body_size)
                        last_activity_paragraph = None
                    else:
                        p = document.add_paragraph(); p.paragraph_format.space_after = Pt(.25)
                        _set_font(p.add_run(left), body_size, True, family=family)
                        last_activity_paragraph = p
            if last_activity_paragraph is not None:
                last_activity_paragraph.paragraph_format.space_after = Pt(max(1.5, activity_gap * .65))

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    return target



def render_structured_cv_docx(document_model: ResumeDocument, target: Path, options: dict | None = None, hidden_sections: list[str] | None = None) -> Path:
    """Render a true multi-page CV from verified canonical facts.

    The CV intentionally uses a visible document label and roomier defaults so
    it cannot be mistaken for the one-page Resume template. It preserves all
    supported canonical sections and does not use Preserve Original Resume mode.
    """
    options = dict(options or {})
    options["font_size"] = max(9.5, float(options.get("font_size", 9.5)))
    options["line_spacing"] = max(1.05, float(options.get("line_spacing", 1.05)))
    options["section_spacing_before"] = max(7.0, float(options.get("section_spacing_before", 7.0)))
    options["section_spacing_after"] = max(2.8, float(options.get("section_spacing_after", 2.8)))
    options["margin_top"] = max(.65, float(options.get("margin_top", .65)))
    options["margin_bottom"] = max(.65, float(options.get("margin_bottom", .65)))
    options["margin_left"] = max(.68, float(options.get("margin_left", .68)))
    options["margin_right"] = max(.68, float(options.get("margin_right", .68)))
    return render_structured_resume_docx(document_model, target, options, hidden_sections, document_label="CURRICULUM VITAE")


def render_cv_letter_docx(letter: str, target: Path, candidate: dict, company: str, title: str, job_location: str = "", date_text: str = "", options: dict | None = None) -> Path:
    """Render the product's CV as a conventional one-page application letter."""
    options = options or {}
    family = "Times New Roman" if str(options.get("font_family", "")).casefold().startswith("times") else "Arial"
    body_size = max(9.5, min(12.0, float(options.get("font_size", 10.5))))
    margin = {"Narrow": .65, "Standard": .8, "Comfortable": .9}.get(str(options.get("margins", "Standard")), .8)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(margin); section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin); section.right_margin = Inches(margin)
    normal = document.styles["Normal"]
    normal.font.name = family; normal._element.rPr.rFonts.set(qn("w:ascii"), family); normal._element.rPr.rFonts.set(qn("w:hAnsi"), family)
    normal.font.size = Pt(body_size); normal.paragraph_format.space_after = Pt(8); normal.paragraph_format.line_spacing = 1.08

    full_name = str(candidate.get("full_name") or "Candidate").strip()
    header = document.add_paragraph(); header.paragraph_format.space_after = Pt(2)
    _set_font(header.add_run(full_name), body_size + 5, True, ACCENT, family=family)
    contact_values = [candidate.get("location"), candidate.get("phone"), candidate.get("email"), candidate.get("linkedin")]
    contact = " | ".join(str(value).strip() for value in contact_values if str(value or "").strip())
    if contact:
        p = document.add_paragraph(); p.paragraph_format.space_after = Pt(12); _set_font(p.add_run(contact), max(8.5, body_size - 1), family=family)
    if date_text:
        p = document.add_paragraph(); p.paragraph_format.space_after = Pt(10); _set_font(p.add_run(date_text), body_size, family=family)
    recipient_lines = [company, job_location]
    for value in recipient_lines:
        if str(value or "").strip():
            p = document.add_paragraph(); p.paragraph_format.space_after = Pt(0); _set_font(p.add_run(str(value).strip()), body_size, family=family)
    subject = document.add_paragraph(); subject.paragraph_format.space_before = Pt(10); subject.paragraph_format.space_after = Pt(10)
    _set_font(subject.add_run(f"Re: {title}"), body_size, True, family=family)

    for raw in re.split(r"\n\s*\n", str(letter or "").strip()):
        lines = [line.strip() for line in raw.splitlines() if line.strip()]
        if not lines:
            continue
        p = document.add_paragraph(); p.paragraph_format.space_after = Pt(8)
        for index, line in enumerate(lines):
            run = p.add_run(line); _set_font(run, body_size, family=family)
            if index < len(lines) - 1:
                run.add_break()
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    return target

def render_professional_one_page(content: str, target: Path, options: dict | None = None) -> Path:
    """Legacy text renderer for non-structured callers."""
    options = options or {}
    family = "Times New Roman" if str(options.get("font_family", "")).casefold().startswith("times") else "Arial"
    body_size = max(8.0, min(12.0, float(options.get("font_size", 9.0))))
    line_spacing = max(.9, min(1.4, float(options.get("line_spacing", 1.0))))
    margins = {"Narrow": .50, "Standard": .68, "Comfortable": .82}
    margin = margins.get(str(options.get("margins", "Standard")), .68)
    section_before = float(options.get("section_spacing_before", {"Compact": 4.0, "Normal": 6.0, "Spacious": 9.0}.get(str(options.get("section_spacing", "Normal")), 6.0)))
    section_after = float(options.get("section_spacing_after", max(2.0, section_before / 2.4)))
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(float(options.get("margin_top", margin)))
    section.bottom_margin = Inches(float(options.get("margin_bottom", margin)))
    section.left_margin = Inches(float(options.get("margin_left", margin)))
    section.right_margin = Inches(float(options.get("margin_right", margin)))
    normal = document.styles["Normal"]
    normal.font.name = family
    normal._element.rPr.rFonts.set(qn("w:ascii"), family)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), family)
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = line_spacing
    raw = [line.rstrip() for line in content.replace("\r", "").split("\n")]
    first = next((line.strip() for line in raw if line.strip()), "Resume")
    used_first = False
    for raw_line in raw:
        line = raw_line.strip()
        if not line:
            continue
        if not used_first and line == first:
            p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(line), float(options.get("name_font_size", body_size + 8.5)), True, ACCENT, family=family); used_first = True; continue
        if _heading(line):
            p = document.add_paragraph(); p.paragraph_format.space_before = Pt(section_before); p.paragraph_format.space_after = Pt(section_after)
            _set_font(p.add_run(line.lstrip("#").strip().rstrip(":")), float(options.get("section_font_size", body_size + 1.1)), True, ACCENT, family=family); _rule(p); continue
        is_bullet = bool(BULLET.match(line)); text = BULLET.sub("", line) if is_bullet else line
        p = document.add_paragraph(); p.paragraph_format.line_spacing = line_spacing; p.paragraph_format.space_after = Pt(float(options.get("paragraph_spacing", 1.1)))
        if is_bullet:
            p.paragraph_format.left_indent = Inches(0.18); p.paragraph_format.first_line_indent = Inches(-0.13)
            _set_font(p.add_run("• "), body_size, True, ACCENT, family=family)
        _set_font(p.add_run(text), body_size, family=family)
    target.parent.mkdir(parents=True, exist_ok=True); document.save(target); return target


def render_resume_docx(content: str, target: Path, options: dict | None = None) -> tuple[Path, str]:
    # Kept as a compatibility wrapper for older callers. Source-based in-place
    # editing is deliberately retired; every Resume uses the controlled renderer.
    return render_professional_one_page(content, target, options or {}), "Professional one-page"
