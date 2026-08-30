from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.resume_ingestion import extract_docx_blocks, normalized_source_text, parse_docx_resume
from app.resume_integrity import ResumeIntegrityValidator
from app.resume_models import EducationItem, PersonalInfo, ResumeDocument


class ResumeIngestionTests(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory(); self.path = Path(self.tmp.name) / "fixture.docx"

    def tearDown(self): self.tmp.cleanup()

    def _fixture(self):
        document = Document(); document.add_paragraph("XUAN NIE"); document.add_paragraph("Ottawa, ON | xuan@example.com")
        document.add_paragraph("EDUCATION").runs[0].bold = True
        row = document.add_table(rows=1, cols=2).rows[0]; row.cells[0].text = "Bachelor of Mechanical and Aerospace Engineering\nCarleton University - Concentration: Aerodynamics"; row.cells[1].text = "Expected Apr 2027\nOttawa, ON"
        document.add_paragraph("TECHNICAL SKILLS").runs[0].bold = True
        document.add_paragraph("Engineering & Analysis: ANSYS (CFX, TurboGrid), FEA, CFD")
        document.add_paragraph("Languages: English")
        document.add_paragraph("ENGINEERING PROJECTS").runs[0].bold = True
        row = document.add_table(rows=1, cols=2).rows[0]; row.cells[0].text = "Aerospace Design Project | AIAA Mission"; row.cells[1].text = "Jan 2023 - Apr 2023"
        document.add_paragraph("• Built a test vehicle")
        document.add_paragraph("ENGINEERING TEAMS & COMPETITIONS").runs[0].bold = True
        document.add_paragraph("Student Design Competition")
        document.save(self.path)

    def test_interleaved_blocks_keep_document_order(self):
        self._fixture(); blocks = extract_docx_blocks(self.path); text = normalized_source_text(blocks)
        self.assertLess(text.index("EDUCATION"), text.index("Bachelor")); self.assertLess(text.index("TECHNICAL SKILLS"), text.index("ENGINEERING PROJECTS"))
        self.assertEqual([block.order for block in blocks], list(range(len(blocks))))

    def test_problem_resume_semantics_are_separate_and_traceable(self):
        self._fixture(); document, blocks = parse_docx_resume(self.path)
        self.assertEqual(document.personal_info.full_name, "XUAN NIE")
        self.assertEqual(document.education[0].institution, "Carleton University")
        self.assertEqual(document.education[0].degree, "Bachelor of Mechanical and Aerospace Engineering")
        self.assertEqual(document.education[0].date_text, "Expected Apr 2027")
        self.assertEqual(document.skills[0].items[0], "ANSYS (CFX, TurboGrid)")
        self.assertEqual(document.languages[0].language, "English")
        self.assertEqual(document.projects[0].name, "Aerospace Design Project")
        self.assertEqual(len(document.projects[0].bullets), 1)
        self.assertTrue(document.projects[0].source_refs); self.assertTrue(document.source_block_ids)
        self.assertFalse(document.unresolved_items); self.assertEqual(ResumeIntegrityValidator().report(document)["status"], "PASS")

    def test_mocked_ai_can_resolve_only_source_backed_unresolved_education(self):
        document = Document(); document.add_paragraph("JANE DOE"); document.add_paragraph("EDUCATION").runs[0].bold = True; document.add_paragraph("B.Eng. | Carleton University") ; document.save(self.path)
        class Resolver:
            def resolve(self, section, blocks):
                return [{"type":"education", "confidence":.95, "institution":"Carleton University", "degree":"B.Eng.", "source_block_ids":[blocks[0].id]}]
        parsed, _ = parse_docx_resume(self.path, Resolver())
        self.assertEqual(parsed.education[0].institution, "Carleton University"); self.assertFalse(parsed.unresolved_items)

    def test_bad_or_unavailable_ai_keeps_content_unresolved(self):
        document = Document(); document.add_paragraph("JANE DOE"); document.add_paragraph("EDUCATION").runs[0].bold = True; document.add_paragraph("B.Eng. | Carleton University"); document.save(self.path)
        class BadResolver:
            def resolve(self, *_): return [{"type":"education", "confidence":.99, "institution":"Invented University", "degree":"B.Eng.", "source_block_ids":["wrong"]}]
        parsed, _ = parse_docx_resume(self.path, BadResolver()); self.assertTrue(parsed.unresolved_items); self.assertEqual(ResumeIntegrityValidator().report(parsed)["status"], "REVIEW_REQUIRED")
        class OfflineResolver:
            def resolve(self, *_): raise OSError("offline")
        parsed, _ = parse_docx_resume(self.path, OfflineResolver()); self.assertTrue(parsed.unresolved_items)

    def test_semantic_validator_rejects_historical_education_merge(self):
        document = ResumeDocument(personal_info=PersonalInfo(full_name="XUAN NIE"), source_block_ids=["block-1"], education=[EducationItem("education-1", institution="Bachelor of Mechanical and Aerospace Engineering / Carleton University", degree="Bachelor of Mechanical and Aerospace Engineering / Carleton University", source_refs=["block-1"])])
        report = ResumeIntegrityValidator().report(document)
        self.assertEqual(report["status"], "REVIEW_REQUIRED")
        self.assertIn("EDUCATION_DEGREE_IN_INSTITUTION", {issue["code"] for issue in report["issues"]})
