from __future__ import annotations

import sqlite3
import tempfile
import unittest
import gc
from pathlib import Path
from types import SimpleNamespace

from app.database import Database
from app.resume_models import ResumeDocument, ResumeLayout, ResumeStyle, resume_document_from_dict, resume_document_to_dict, resume_layout_from_dict, resume_layout_to_dict, resume_style_from_dict, resume_style_to_dict
from app.resume_parser import parse_resume_text, resume_document_to_text
from app.resume_integrity import ResumeIntegrityValidator, apply_resume_operations, apply_resume_operations_best_effort
from app.services import ResumeService
from app.pdf_export import render_resume_document_pdf
from app.docx_export import render_resume_docx


class ResumeV2Tests(unittest.TestCase):
    def setUp(self): self.tmp = tempfile.TemporaryDirectory()
    def tearDown(self): self.tmp.cleanup()

    def test_models_round_trip(self):
        document = ResumeDocument(source_text="Jane Doe\nEMAIL: jane@example.com")
        self.assertEqual(resume_document_from_dict(resume_document_to_dict(document)).source_text, document.source_text)
        self.assertEqual(resume_style_from_dict(resume_style_to_dict(ResumeStyle(font_family="Times"))).font_family, "Times")
        self.assertEqual(resume_layout_from_dict(resume_layout_to_dict(ResumeLayout(hidden_sections=["projects"]))).hidden_sections, ["projects"])

    def test_parser_preserves_known_and_unknown_lines(self):
        source = "Jane Doe\njane@example.com | Ottawa\nSUMMARY\nMechanical engineer\nSKILLS\nCAD, Python\nPROJECTS\n- Built a test rig\nAWARDS\nDean's list"
        document = parse_resume_text(source); rendered = resume_document_to_text(document)
        for line in ("Jane Doe", "CAD, Python", "Built a test rig", "Dean's list"):
            self.assertIn(line, rendered)
        self.assertNotIn("Mechanical engineer", rendered)
        self.assertIn("Mechanical engineer", document.source_text)
        self.assertEqual(document.summary, "")
        self.assertEqual(document.personal_info.email, "jane@example.com")
        self.assertEqual(rendered.count("jane@example.com"), 1)

    def test_uppercase_name_is_personal_info_not_a_custom_section(self):
        document = parse_resume_text("XUAN NIE\nOttawa, ON | xuan@example.com\nEDUCATION\nCarleton University")
        self.assertEqual(document.personal_info.full_name, "XUAN NIE")
        self.assertFalse(document.custom_sections)

    def test_database_migrates_legacy_resume_versions(self):
        path = Path(self.tmp.name) / "legacy.db"
        with sqlite3.connect(path) as db:
            db.execute("CREATE TABLE resume_versions (id INTEGER PRIMARY KEY, job_id INTEGER, version_name TEXT NOT NULL, source_path TEXT, content TEXT NOT NULL, changes_json TEXT NOT NULL DEFAULT '[]', model_used TEXT, created_at TEXT NOT NULL, approved INTEGER NOT NULL DEFAULT 0, rejected INTEGER NOT NULL DEFAULT 0)")
        database = Database(path)
        with database.connection() as db:
            columns = {row[1] for row in db.execute("PRAGMA table_info(resume_versions)")}
        self.assertTrue({"document_json", "style_json", "layout_json"}.issubset(columns))
        del database; gc.collect()

    def test_new_version_keeps_structured_json(self):
        database = Database(Path(self.tmp.name) / "new.db")
        version_id = database.add_resume_version(version_name="resume_v001", content="Jane Doe", document_json='{"schema_version":1}', style_json='{"template_id":"modern"}', layout_json='{"layout_mode":"single_column"}')
        version = next(value for value in database.resume_versions() if value["id"] == version_id)
        self.assertEqual(version["document_json"], '{"schema_version":1}')


    def test_imported_source_path_uses_import_timestamp_not_preserved_mtime(self):
        import os
        root = Path(self.tmp.name) / "source-order"
        root.mkdir()
        legacy_pdf = root / "source-20260828-120000-000001.pdf"
        new_docx = root / "source-20260828-220000-000001.docx"
        legacy_pdf.write_bytes(b"legacy")
        new_docx.write_bytes(b"new")
        # Simulate copy2 preserving an old mtime on the newly imported DOCX.
        os.utime(legacy_pdf, (2000000000, 2000000000))
        os.utime(new_docx, (1000000000, 1000000000))
        service = ResumeService(Database(Path(self.tmp.name) / "source-order.db"), object())
        service.paths = SimpleNamespace(resumes_original=root)
        self.assertEqual(service.imported_source_path(), new_docx)

    def test_new_resume_import_rejects_non_docx(self):
        source = Path(self.tmp.name) / "resume.pdf"; source.write_bytes(b"not a resume")
        service = ResumeService(Database(Path(self.tmp.name) / "import.db"), object())
        with self.assertRaisesRegex(ValueError, "editable DOCX"):
            service.import_original(str(source))

    def test_empty_or_corrupt_docx_is_rejected_without_retaining_backup(self):
        from docx import Document
        original_dir = Path(self.tmp.name) / "original"; original_dir.mkdir()
        source = Path(self.tmp.name) / "empty.DOCX"; Document().save(source)
        service = ResumeService(Database(Path(self.tmp.name) / "empty.db"), object())
        service.paths = SimpleNamespace(resumes_original=original_dir)
        with self.assertRaisesRegex(ValueError, "could not read"):
            service.import_original(str(source))
        self.assertEqual(list(original_dir.glob("source-*")), [])

    def test_structured_pdf_renderer_handles_long_bullets_and_hidden_sections(self):
        document = parse_resume_text("Jane Doe\nPROJECTS\n- " + "Detailed engineering result " * 45 + "\nSKILLS\nPython, CAD")
        layout = ResumeLayout(hidden_sections=["skills"])
        rendered = resume_document_to_text(document, layout.hidden_sections)
        self.assertNotIn("Python, CAD", rendered)
        target = Path(self.tmp.name) / "resume.pdf"
        render_resume_document_pdf(document, target, ResumeStyle(), layout)
        self.assertGreater(target.stat().st_size, 500)

    def test_service_loads_legacy_content_as_document(self):
        database = Database(Path(self.tmp.name) / "fallback.db")
        version_id = database.add_resume_version(version_name="legacy", content="Jane Doe\nSKILLS\nPython")
        version = next(value for value in database.resume_versions() if value["id"] == version_id)
        service = ResumeService(database, object())
        self.assertEqual(service.document_from_version(version).personal_info.full_name, "Jane Doe")

    def test_structured_patch_keeps_section_ownership_and_rejects_section_text(self):
        document = parse_resume_text("Jane Doe\nEDUCATION\nCarleton University\nSKILLS\nPython, CAD\nPROJECTS\n- Built a test rig\nEXPERIENCE\n- Tested hardware")
        project = document.projects[0]; original_education = list(document.education[0].extra_lines); original_skills = list(document.skills[0].items)
        edited, changes = apply_resume_operations(document, [{"operation": "replace_bullet", "entity_type": "project", "entity_id": project.id, "bullet_id": project.bullets[0].id, "value": "Built and tested a student test rig", "reason": "clarity"}])
        self.assertEqual(edited.education[0].extra_lines, original_education)
        self.assertEqual(edited.skills[0].items, original_skills)
        self.assertEqual(edited.projects[0].bullets[0].text, "Built and tested a student test rig")
        self.assertEqual(len(changes), 1)
        with self.assertRaisesRegex(ValueError, "suspicious structural"):
            apply_resume_operations(document, [{"operation": "replace_bullet", "entity_type": "project", "entity_id": project.id, "bullet_id": project.bullets[0].id, "value": "EDUCATION\nWrong place", "reason": "bad"}])

    def test_summary_patch_is_rejected_after_summary_retirement(self):
        document = parse_resume_text("Jane Doe\nPROJECTS\n- Built a test rig")
        with self.assertRaisesRegex(ValueError, "Unsupported resume AI operation"):
            apply_resume_operations(document, [{
                "operation": "replace_summary",
                "entity_type": "summary",
                "entity_id": "summary",
                "value": "Mechanical engineering student with hands-on design experience.",
            }])

    def test_best_effort_patch_skips_one_bad_operation_and_keeps_valid_edits(self):
        document = parse_resume_text("Jane Doe\nPROJECTS\n- Built a test rig\nSKILLS\nPython, CAD")
        project = document.projects[0]
        edited, changes, rejected = apply_resume_operations_best_effort(document, [
            {"operation": "replace_bullet", "entity_type": "project", "entity_id": project.id, "bullet_id": project.bullets[0].id, "value": "Built and tested a student test rig"},
            {"operation": "replace_bullet", "entity_type": "custom", "entity_id": "section-999", "bullet_id": "bullet-999", "value": "Unsupported edit"},
        ])
        self.assertEqual(edited.projects[0].bullets[0].text, "Built and tested a student test rig")
        self.assertEqual(len(changes), 1)
        self.assertEqual(len(rejected), 1)
        self.assertIn("unknown or wrong entity", rejected[0]["reason"])

    def test_skill_patch_does_not_require_value_field(self):
        document = parse_resume_text("Jane Doe\nSKILLS\nEngineering & Analysis: Python, MATLAB")
        skill = document.skills[0]
        edited, changes = apply_resume_operations(document, [{
            "operation": "replace_skill_items",
            "entity_type": "skills",
            "entity_id": skill.id,
            "items": ["Python", "MATLAB", "CFD"],
            "reason": "keyword alignment",
        }])
        self.assertEqual(edited.skills[0].items, ["Python", "MATLAB", "CFD"])
        self.assertEqual(len(changes), 1)

    def test_cv_rejects_resume_style_ai_output(self):
        class BadAI:
            def generate_json(self, *_args, **_kwargs):
                return {"operations": [{"operation": "replace_bullet", "entity_type": "custom", "entity_id": "section-23", "bullet_id": "missing", "value": "bad"}]}, "fake"
        root = Path(self.tmp.name) / "cv-data"
        paths = SimpleNamespace(resumes_original=root / "original", resumes_generated=root / "generated", resumes_approved=root / "approved", cache=root / "cache")
        for folder in vars(paths).values(): folder.mkdir(parents=True, exist_ok=True)
        database = Database(Path(self.tmp.name) / "cv-safe.db")
        job_id, _ = database.upsert_job({"company": "Example", "title": "Engineer", "location": "Ottawa", "source": "test", "url": "https://example.test/cv-job", "description": "CAD"})
        (paths.resumes_original / "original-fixture.txt").write_text("Jane Doe\nPROJECTS\n- Built a test rig", encoding="utf-8")
        service = ResumeService(database, BadAI()); service.paths = paths
        service.render_version_outputs = lambda version: (paths.resumes_generated / "safe.pdf", paths.resumes_generated / "safe.docx")
        with self.assertRaisesRegex(ValueError, "CV letter"):
            service.optimize(database.job(job_id), document_type="CV")
        self.assertEqual(database.resume_versions(), [])

    def test_cv_generation_persists_a_letter_instead_of_resume_json(self):
        class LetterAI:
            request = ""
            def generate_json(self, *_args, **_kwargs):
                self.request = _args[2]
                letter = (
                    "Dear Hiring Manager,\n\n"
                    "I am applying for the Engineer opportunity at Example. My background includes hands-on engineering work and a project in which I Built a test rig. I am interested in contributing that practical approach to the responsibilities described for this role.\n\n"
                    "The project required me to turn an engineering need into a working physical system. That experience strengthened my ability to organize technical work, examine results carefully, and communicate design decisions in a clear and useful way. These are the habits I would bring to the team while learning its specific processes and tools.\n\n"
                    "I am particularly interested in the opportunity to support collaborative engineering work at Example. I would welcome the chance to discuss how my verified project experience, careful working style, and interest in the position could support the team without overstating qualifications that are not in my background.\n\n"
                    "Thank you for considering my application. I would be pleased to discuss the role and my experience further.\n\n"
                    "Sincerely,\nJane Doe"
                )
                return {"letter": letter, "facts_used": [{"fact": "Test rig construction experience", "source": "resume"}], "warnings": []}, "fake"
        root = Path(self.tmp.name) / "cv-letter-data"
        paths = SimpleNamespace(resumes_original=root / "original", resumes_generated=root / "generated", resumes_approved=root / "approved", cache=root / "cache")
        for folder in vars(paths).values(): folder.mkdir(parents=True, exist_ok=True)
        database = Database(Path(self.tmp.name) / "cv-letter.db")
        job_id, _ = database.upsert_job({"company": "Example", "title": "Engineer", "location": "Ottawa", "source": "test", "url": "https://example.test/cv-letter", "description": "Collaborative engineering role"})
        (paths.resumes_original / "original-fixture.txt").write_text("Jane Doe\nPROJECTS\n- Built a test rig", encoding="utf-8")
        ai = LetterAI(); service = ResumeService(database, ai); service.paths = paths
        service.render_version_outputs = lambda version: (paths.resumes_generated / "letter.pdf", paths.resumes_generated / "letter.docx")
        from unittest.mock import patch
        with patch("app.services.load_settings", return_value={"profile": {"additional_facts": "Verified design lab experience"}, "resume_pdf": {}, "generation_prompts": {}}):
            result = service.optimize(database.job(job_id), document_type="CV")
        version = database.resume_versions()[0]
        self.assertEqual(result["document_type"], "CV")
        self.assertTrue(version["content"].startswith("Dear Hiring Manager,"))
        self.assertEqual(version["template_version"], "cv-letter-v1")
        self.assertEqual(__import__("json").loads(version["document_json"])["kind"], "cv_letter")
        self.assertNotIn('"education"', version["document_json"])
        self.assertIn("Verified design lab experience", ai.request)
        self.assertIn("Collaborative engineering role", ai.request)

    def test_cv_fact_grounding_accepts_a_traceable_paraphrase(self):
        source = "Designed and tested a rocket-propelled wind tunnel using sensor data acquisition."
        self.assertTrue(ResumeService._cv_fact_is_grounded("Wind tunnel design and testing experience", source))

    def test_cv_fact_grounding_rejects_an_unsupported_claim(self):
        source = "Built and tested a wind tunnel for a university project."
        self.assertFalse(ResumeService._cv_fact_is_grounded("Managed aerospace certification programs and supplier contracts", source))

    def test_regeneration_uses_the_drafts_original_resume_snapshot(self):
        class RecordingAI:
            request = ""
            def generate_json(self, *_args, **_kwargs):
                self.request = _args[2]
                return {"operations": []}, "fake"
        root = Path(self.tmp.name) / "snapshot-data"
        paths = SimpleNamespace(resumes_original=root / "original", resumes_generated=root / "generated", resumes_approved=root / "approved", cache=root / "cache")
        for folder in vars(paths).values(): folder.mkdir(parents=True, exist_ok=True)
        old = paths.resumes_original / "original-20260101-000000.txt"
        current = paths.resumes_original / "original-20260201-000000.txt"
        old.write_text("Jane Doe\nPROJECTS\n- OLD SNAPSHOT PROJECT", encoding="utf-8")
        current.write_text("Jane Doe\nPROJECTS\n- NEW CURRENT PROJECT", encoding="utf-8")
        database = Database(Path(self.tmp.name) / "snapshot.db")
        job_id, _ = database.upsert_job({"company":"Example", "title":"Engineer", "location":"Ottawa", "url":"https://example.test/snapshot", "description":"Engineering", "description_hash":"snapshot"})
        version_id = database.add_resume_version(job_id=job_id, version_name="resume_v001", content="old draft", source_path=str(old), document_type="Resume")
        existing = next(value for value in database.resume_versions() if value["id"] == version_id)
        ai = RecordingAI(); service = ResumeService(database, ai); service.paths = paths
        service.render_version_outputs = lambda version: (paths.resumes_generated / "safe.pdf", paths.resumes_generated / "safe.docx")
        from unittest.mock import patch
        with patch("app.services.load_settings", return_value={"profile": {}, "resume_pdf": {}, "generation_prompts": {}}):
            service.optimize(database.job(job_id), document_type="Resume", replace_version=existing)
        self.assertIn("OLD SNAPSHOT PROJECT", ai.request)
        self.assertNotIn("NEW CURRENT PROJECT", ai.request)
        updated = next(value for value in database.resume_versions() if value["id"] == version_id)
        self.assertEqual(updated["source_path"], str(old))

    def test_invalid_ai_patch_never_persists_a_draft(self):
        class BadAI:
            def generate_json(self, *_args, **_kwargs):
                return {"operations": [{"operation": "replace_bullet", "entity_type": "education", "entity_id": "education-1", "bullet_id": "missing", "value": "bad"}]}, "fake"
        root = Path(self.tmp.name) / "data"
        paths = SimpleNamespace(resumes_original=root / "original", resumes_generated=root / "generated", resumes_approved=root / "approved", cache=root / "cache")
        for folder in vars(paths).values(): folder.mkdir(parents=True, exist_ok=True)
        database = Database(Path(self.tmp.name) / "safe.db")
        job_id, _ = database.upsert_job({"company": "Example", "title": "Engineer", "location": "Ottawa", "source": "test", "url": "https://example.test/job", "description": "CAD"})
        (paths.resumes_original / "original-fixture.txt").write_text("Jane Doe\nPROJECTS\n- Built a test rig", encoding="utf-8")
        service = ResumeService(database, BadAI()); service.paths = paths
        with self.assertRaisesRegex(ValueError, "unknown or wrong entity"):
            service.optimize(database.job(job_id))
        self.assertEqual(database.resume_versions(), [])

    def test_docx_and_pdf_use_identical_section_order_and_content(self):
        from docx import Document
        from pypdf import PdfReader
        document = parse_resume_text("Jane Doe\njane@example.com | Ottawa\nEDUCATION\nCarleton University\nSKILLS\nPython, CAD\nPROJECTS\n- Built a test rig\nEXPERIENCE\n- Tested hardware")
        text = resume_document_to_text(document)
        docx = Path(self.tmp.name) / "resume.docx"; pdf = Path(self.tmp.name) / "resume.pdf"
        render_resume_docx(text, docx, {"font_family": "Arial", "font_size": 9, "line_spacing": 1.0, "margins": "Standard"})
        render_resume_document_pdf(document, pdf, ResumeStyle(), ResumeLayout())
        docx_text = "\n".join(paragraph.text for paragraph in Document(docx).paragraphs)
        pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(str(pdf)).pages)
        labels = ["EDUCATION", "SKILLS", "PROJECTS", "EXPERIENCE"]
        self.assertEqual([label for label in labels if label in docx_text], labels)
        self.assertEqual([label for label in labels if label in pdf_text], labels)
        self.assertIn("Built a test rig", docx_text); self.assertIn("Built a test rig", pdf_text)

class StructuredDocxLayoutTests(unittest.TestCase):
    def test_structured_renderer_keeps_semantic_layout_and_dedupes_phone(self):
        from tempfile import TemporaryDirectory
        from docx import Document
        from app.docx_export import render_structured_resume_docx
        from app.resume_models import ResumeDocument, PersonalInfo, EducationItem, ProjectItem, ResumeBullet, SkillGroup, LanguageItem
        with TemporaryDirectory() as tmp:
            doc = ResumeDocument(
                personal_info=PersonalInfo(full_name="Jane Doe", email="jane@example.com", phone="587) 892-4643", location="Ottawa, ON", other_lines=["(587) 892-4643"]),
                education=[EducationItem(id="education-1", institution="Carleton University", degree="Bachelor of Engineering", concentration="Aerospace", location="Ottawa, ON", date_text="Expected Apr 2027")],
                projects=[ProjectItem(id="project-1", name="Aero Project", subtitle="Wind Tunnel", date_text="Sep 2024 – Apr 2025", bullets=[ResumeBullet(id="bullet-1", text="Built and tested the system.")])],
                skills=[SkillGroup(id="skills-1", name="Engineering & Analysis", items=["ANSYS (CFX, TurboGrid)", "MATLAB"])],
                languages=[LanguageItem(id="language-1", language="English")],
                section_order=["education", "skills", "languages", "projects"],
            )
            target = Path(tmp) / "resume.docx"
            render_structured_resume_docx(doc, target, {"font_size": 9})
            parsed = Document(target)
            text = "\n".join(p.text for p in parsed.paragraphs) + "\n" + "\n".join(cell.text for table in parsed.tables for row in table.rows for cell in row.cells)
            self.assertIn("TECHNICAL SKILLS", text)
            self.assertIn("ENGINEERING PROJECTS", text)
            self.assertIn("Bachelor of Engineering", text)
            self.assertIn("Expected Apr 2027", text)
            self.assertIn("Aero Project", text)
            self.assertEqual(text.count("892-4643"), 1)

class CvRendererTests(unittest.TestCase):
    def test_cv_docx_uses_application_letter_layout(self):
        from tempfile import TemporaryDirectory
        from pathlib import Path
        from docx import Document
        from pypdf import PdfReader
        from app.docx_export import render_cv_letter_docx
        from app.pdf_export import render_cv_letter_pdf
        with TemporaryDirectory() as directory:
            target = Path(directory) / "test-cv-letter.docx"
            render_cv_letter_docx("Dear Hiring Manager,\n\nI am interested in this role.\n\nSincerely,\nTest Candidate", target, {"full_name": "Test Candidate", "email": "test@example.com"}, "Example Company", "Engineer", "Ottawa", "August 29, 2026", {})
            text = "\n".join(p.text for p in Document(target).paragraphs)
            self.assertIn("Test Candidate", text)
            self.assertIn("Dear Hiring Manager", text)
            self.assertIn("Re: Engineer", text)
            self.assertNotIn("EDUCATION", text)
            pdf = Path(directory) / "test-cv-letter.pdf"
            render_cv_letter_pdf("Dear Hiring Manager,\n\nI am interested in this role.\n\nSincerely,\nTest Candidate", pdf, {"full_name": "Test Candidate", "email": "test@example.com"}, "Example Company", "Engineer", "Ottawa", "August 29, 2026", {})
            pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(str(pdf)).pages)
            self.assertIn("Test Candidate", pdf_text)
            self.assertIn("Dear Hiring Manager", pdf_text)
            self.assertIn("Re: Engineer", pdf_text)
            self.assertNotIn("EDUCATION", pdf_text)
